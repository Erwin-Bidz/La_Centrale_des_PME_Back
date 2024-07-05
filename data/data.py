import json
from collections import defaultdict
from docx import Document
import os

def extract_tables_from_docx(file_path):
    doc = Document(file_path)
    tables = []
    current_title = None
    
    for paragraph in doc.paragraphs:
        # Detect table titles
        if "Tableau" in paragraph.text:
            current_title = paragraph.text.split(" : ", 1)[1].strip()
        else:
            current_title = None
        
        # Capture tables with the most recent title
        if current_title:
            for table in doc.tables:
                data = []
                keys = None
                for i, row in enumerate(table.rows):
                    cells = row.cells
                    text_cells = [cell.text.strip() for cell in cells]
                    if i == 0:
                        keys = text_cells
                    else:
                        if len(text_cells) == len(keys):
                            data.append({keys[j]: text_cells[j] for j in range(len(keys))})
                tables.append((current_title, data))
                
    return tables

def group_tables_by_structure(tables):
    grouped = defaultdict(list)
    for title, table in tables:
        grouped[title].append(table)
    return grouped

def convert_to_json(grouped_tables):
    # Ensures the JSON output uses UTF-8 encoding without escape sequences for non-ASCII characters
    return {structure: json.dumps(tables, ensure_ascii=False, indent=4) for structure, tables in grouped_tables.items()}

def save_json_files(grouped_json):
    for structure, json_data in grouped_json.items():
        filename = f"{structure.replace(' ', '_').replace('/', '_').replace(':', '_')}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(json_data)

if __name__ == "__main__":
    file_path = 'C:/Users/user/Downloads/Telegram Desktop/TROISIEME3_LIVRABLE_RAPPORT_CARTOGRAPHIE_DES_MIGRANTS_CAMEROUN.docx'
    
    tables = extract_tables_from_docx(file_path)
    grouped_tables = group_tables_by_structure(tables)
    grouped_json = convert_to_json(grouped_tables)
    save_json_files(grouped_json)

    print("Tables extracted and saved as JSON files in the current directory")







#file_path = r"C:/Users/user/Downloads/Telegram Desktop/TROISIEME2_LIVRABLE_RAPPORT_CARTOGRAPHIE_DES_MIGRANTS_CAMEROUN.docx"
